# generate_report.py — MediaCraft AI 报表生成器
import os, json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set table cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        set_cell_shading(cell, '1a1a2e')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if r % 2 == 0:
                set_cell_shading(cell, 'f0f0f5')
            if col_widths and c < len(col_widths):
                cell.width = Cm(col_widths[c])

    doc.add_paragraph('')
    return table

def build_report():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== TITLE PAGE =====
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MediaCraft Music')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(212, 168, 67)
    run.bold = True
    run.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('YouTube 运营策略手册')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(26, 26, 46)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}  |  版本 v1.0  |  机密')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ===== STYLE SETUP =====
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ===== 1. EXECUTIVE SUMMARY =====
    doc.add_heading('一、执行摘要', level=1)
    doc.add_paragraph(
        'MediaCraft Music 是基于 AI 音乐生成技术的免版税音乐素材品牌。'
        '本手册制定了 YouTube 频道的完整运营策略，包括：品牌定位、内容矩阵、'
        'SEO 优化、发布节奏、变现路径。目标是在 3-6 个月内建立稳定的搜索流量'
        '和订阅增长，同时通过音乐授权实现现金流。'
    )

    # ===== 2. CHANNEL POSITIONING =====
    doc.add_heading('二、频道定位', level=1)
    add_styled_table(doc,
        ['项目', '内容'],
        [
            ['频道名', 'MediaCraft Music'],
            ['本质', '免版税 AI 音乐素材库 — 既是 YouTube 频道，也是产品目录'],
            ['一句卖点', 'Royalty-free cinematic music for creators,\nfilmmakers, and game developers'],
            ['对标频道', 'Mokka Music (75K) / Scott Buckley (156K) / Whitesand (121K)'],
        ],
        col_widths=[4, 12]
    )

    # ===== 3. MUSIC TYPE MATRIX =====
    doc.add_heading('三、音乐类型矩阵（打什么牌）', level=1)
    add_styled_table(doc,
        ['类型', '占比', '搜索意图', '典型标题', 'Suno 提示词方向'],
        [
            ['🎬 史诗管弦\nEpic Orchestral', '35%', 'epic music for\ntrailer', 'Heroic Epic Orchestral |\nCinematic Battle Music', 'epic cinematic, brass\nfanfare, sweeping strings'],
            ['🎥 电影背景\nCinematic BG', '25%', 'background music\nfor documentary', 'Documentary Background |\nCinematic Instrumental', 'cinematic background,\nsubtle strings, pad'],
            ['🌑 暗黑预告\nDark/Trailer', '20%', 'dark trailer\nmusic', 'Dark Cinematic Trailer |\nEpic Hybrid Orchestral', 'dark orchestral hybrid,\ntension, bass drops'],
            ['🏰 奇幻冒险\nFantasy', '10%', 'fantasy music\nfor games', 'Fantasy Adventure |\nOrchestral Soundtrack', 'fantasy orchestral,\nchoir, magical'],
            ['🌌 氛围太空\nAmbient', '10%', 'ambient\nbackground music', 'Atmospheric Space Ambient |\nCinematic Drone', 'ambient cinematic,\ndeep drone, ethereal'],
        ],
        col_widths=[3, 1.5, 2.5, 4, 4]
    )

    # ===== 4. UPLOAD SCHEDULE =====
    doc.add_heading('四、更新频率', level=1)
    add_styled_table(doc,
        ['阶段', '单曲/周', 'Shorts/周', '合集', '社区帖/周'],
        [
            ['第1个月（冷启动）', '3-4 首', '3-5 条', '每2周 1个', '2 次'],
            ['第2-3个月（建库期）', '2-3 首', '3-5 条', '每月 2个', '1-2 次'],
            ['第4个月+（稳定期）', '1-2 首', '2-3 条', '每月 1个', '1 次'],
        ],
        col_widths=[4, 2, 2, 3, 2]
    )
    doc.add_paragraph('发布日推荐：周二/周四（YouTube 观众活跃高峰前 2-3 小时发布）')

    # ===== 5. SINGLE vs COMPILATION =====
    doc.add_heading('五、单曲 vs 合集 — 双轨策略', level=1)
    add_styled_table(doc,
        ['', '单曲 (3-5分钟)', '合集 (1-2小时)'],
        [
            ['核心作用', '🔍 搜索发现引擎', '⏱ 观看时长引擎'],
            ['算法信号', 'CTR + SEO 关键词排名', 'Watch Time + 会话时长'],
            ['听众心态', '"我在找一首特定的歌"', '"我想当背景音放一小时"'],
            ['购买转化', '高（精准搜索 = 购买意图）', '低（被动收听）'],
            ['SEO 策略', '精确关键词，15个标签', '场景 + 曲风组合，合集标题含 "1 HOUR"'],
        ],
        col_widths=[4, 6, 6]
    )
    doc.add_paragraph('规则：每首单曲都出。每 8-12 首按场景打包一个合集。合集的 Chapter 分段标注曲名和关键词。同一个曲目放进多个播放列表。')

    # ===== 6. SEO CHECKLIST =====
    doc.add_heading('六、SEO 执行清单', level=1)

    doc.add_heading('6.1 标题模板', level=2)
    add_styled_table(doc,
        ['场景', '模板', '示例'],
        [
            ['单曲-史诗', '[MOOD] [GENRE] |\n[USE] | Royalty Free', 'Heroic Epic Orchestral |\nCinematic Battle | Royalty Free'],
            ['单曲-场景', '[USE CASE] Background\nMusic | [GENRE]', 'Documentary Background Music |\nCinematic Instrumental'],
            ['合集', '1 HOUR [GENRE] Music for\n[USE] | No Copyright', '1 HOUR Epic Music for Gaming |\nPowerful Orchestral Mix'],
        ],
        col_widths=[3, 5, 8]
    )

    doc.add_heading('6.2 标签策略（前 5 权重最高）', level=2)
    add_styled_table(doc,
        ['位置', '标签内容'],
        [
            ['1-2', '主关键词: epic orchestral music, royalty free orchestral'],
            ['3', '品牌词: MediaCraft Music'],
            ['4-5', '场景词: trailer music, background music'],
            ['6-10', '长尾: cinematic music for videos, film score instrumental, copyright free epic'],
            ['11-15', '补充: heroic soundtrack, orchestral music free, music for creators'],
        ],
        col_widths=[2.5, 13.5]
    )

    doc.add_heading('6.3 购买链接优先级', level=2)
    add_styled_table(doc,
        ['位置', '优先级', '说明'],
        [
            ['Pinned Comment', '🥇 最高', '移动端最显眼，CTR 最高'],
            ['Description 前两行', '🥈', '搜索片段可见，桌面端重要'],
            ['Channel Banner', '🥉', '品牌印象，间接转化'],
        ],
        col_widths=[5, 2.5, 8.5]
    )

    # ===== 7. THUMBNAIL =====
    doc.add_heading('七、缩略图标准', level=1)
    add_styled_table(doc,
        ['参数', '值'],
        [
            ['背景色', '#0D0D0D 深黑'],
            ['强调色（唯一彩色）', '#D4A843 暖金'],
            ['主标题字体', 'Bebas Neue / Anton，白色 + Glow'],
            ['文字规则', '1-3 个关键词，不堆砌'],
            ['底部品牌条', '5% 高度金色色带，所有视频统一'],
            ['CTR 底线', '低于 4% = 立即更换'],
        ],
        col_widths=[5, 11]
    )

    # ===== 8. 30-DAY PLAN =====
    doc.add_heading('八、前 30 天执行表', level=1)
    add_styled_table(doc,
        ['天', '行动', '产出'],
        [
            ['Day 1', '频道装修：横幅+头像+About+预告片', '频道完工上线'],
            ['Day 2', '首发 3 首最强单曲（不同风格测试）', '3 新视频'],
            ['Day 3-4', '1 Short + 1 首新单曲', '每日上新'],
            ['Day 5', '第一组合集（用 Day 2-4 曲目）', '合集上线'],
            ['Day 7', '回复所有评论 + 分析第一周数据', '找到最强风格'],
            ['Day 14', '第 2 组合集 + 3 个播放列表发布', '列表架构就位'],
            ['Day 21', '深度分析 → 复刻最强曲风', '策略验证完毕'],
            ['Day 30', '30 天报告：播放量/CTR/留存/订阅', '定下第 2 月方向'],
        ],
        col_widths=[2, 8, 6]
    )
    doc.add_paragraph('30 天目标：15-20 首单曲 + 3-4 合集 + 15-20 Shorts。订阅 50-200 是正常水平。')

    # ===== 9. PRODUCTION ROADMAP =====
    doc.add_heading('九、批量生产路线图', level=1)
    add_styled_table(doc,
        ['批次', '数量', '类型分配', '完成时间'],
        [
            ['Batch 1', '10 首', '4史诗 + 3电影 + 2暗黑 + 1奇幻', '第 1 周'],
            ['Batch 2', '12 首', '4史诗 + 3电影 + 3暗黑 + 1氛围 + 1奇幻', '第 2-3 周'],
            ['Batch 3+', '15 首/周', '按数据反馈调整比例', '第 4 周起持续'],
        ],
        col_widths=[3, 2.5, 7, 3.5]
    )

    # ===== 10. KPI =====
    doc.add_heading('十、关键数据指标', level=1)
    add_styled_table(doc,
        ['指标', '新手正常值', '优化目标'],
        [
            ['CTR（点击率）', '4-6%', '>8%'],
            ['30 秒留存', '>70%', '>80%'],
            ['3 分钟留存', '>45%', '>55%'],
            ['月播放时长', '100-300h', '500h+'],
            ['周发布频率', '1-2 次', '3-4 次'],
        ],
        col_widths=[4, 4, 4]
    )

    # ===== 11. BENCHMARKS =====
    doc.add_heading('十一、对标频道', level=1)
    add_styled_table(doc,
        ['频道', '订阅', '风格', '频率', '值得学习'],
        [
            ['Mokka Music', '75.4K', '电子/多流派', '每周 2 首', '频率稳定，916 视频的复利效应'],
            ['Scott Buckley', '156K', '电影/管弦', '中低频', '描述栏 CC 协议写得很清楚'],
            ['Silverman Sound', '~20K', '管弦配乐', '低频', '"像好莱坞"的产品定位'],
            ['Savfk Music', '~30K', '史诗管弦', '中频', '合集+播放列表结构优秀'],
        ],
        col_widths=[3, 1.5, 2.5, 2, 6]
    )
    doc.add_paragraph('核心教训：这些频道没有一个靠爆款起来。全部靠 100+ 个视频的长尾复利效应。')

    # ===== FOOTER =====
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— MediaCraft AI · 机密文件 · 仅供内部使用 —')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(160, 160, 160)
    run.italic = True

    # Save
    out_dir = 'd:/自媒体运营/context'
    out_path = os.path.join(out_dir, 'MediaCraft_Music_运营策略手册_v1.docx')
    os.makedirs(out_dir, exist_ok=True)
    doc.save(out_path)
    print(f'Report saved: {out_path}')
    print(f'Size: {os.path.getsize(out_path)/1024:.0f} KB')

if __name__ == '__main__':
    build_report()
